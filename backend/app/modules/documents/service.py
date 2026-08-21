import hashlib
import io
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any, BinaryIO
from uuid import UUID, uuid4

import pypdfium2 as pdfium
from docx import Document as open_docx
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from PIL import Image
from pypdf import PdfReader
from pypdf.errors import PdfReadError
from sqlalchemy import delete, select
from sqlalchemy.orm import Session

from backend.app.integrations.ocr import OcrEngine, PaddleOcrEngine
from backend.app.integrations.storage.local import LocalFileStore
from backend.app.modules.contracts.models import ContractFile, FileObject
from backend.app.modules.documents.models import (
    DocumentBlock,
    DocumentPage,
    DocumentVersion,
    SourceSpan,
)
from backend.app.modules.identity.models import Organization
from backend.app.modules.identity.organization import DEFAULT_PAGE_LIMIT, organization_settings
from backend.app.modules.retention.service import (
    create_file_write_journal,
    finalize_file_write_journal,
)
from backend.app.observability.metrics import observe_ocr_page
from backend.app.shared.db import UnitOfWork

PARSER_VERSION = "2026-08-19-2"
OCR_LOW_CONFIDENCE_DEFAULT = 0.80


class DocumentParseError(RuntimeError):
    def __init__(self, code: str, message: str) -> None:
        self.code = code
        self.message = message
        super().__init__(message)


@dataclass(frozen=True, slots=True)
class ParsedBlock:
    text: str
    block_type: str
    paragraph_no: int | None = None
    table_path: str | None = None
    bbox: dict[str, float] | None = None


@dataclass(frozen=True, slots=True)
class ParsedPage:
    page_no: int
    width: float | None
    height: float | None
    text: str
    blocks: list[ParsedBlock]
    image_bytes: bytes
    ocr_status: str
    ocr_confidence: float | None
    error_code: str | None = None
    error_message: str | None = None


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    parser_name: str
    page_count: int
    ocr_status: str
    pages: list[ParsedPage]
    blocks: list[ParsedBlock]
    text: str


def parse_contract_file(
    session: Session,
    *,
    organization_id: UUID,
    contract_file_id: UUID,
    file_store: LocalFileStore,
    ocr_engine: OcrEngine | None = None,
    ocr_low_confidence_threshold: float | None = None,
) -> DocumentVersion:
    source = session.execute(
        select(ContractFile, FileObject)
        .join(
            FileObject,
            (FileObject.organization_id == ContractFile.organization_id)
            & (FileObject.id == ContractFile.file_object_id),
        )
        .where(
            ContractFile.organization_id == organization_id,
            ContractFile.id == contract_file_id,
        )
    ).one_or_none()
    if source is None:
        raise DocumentParseError("DOCUMENT_FILE_NOT_FOUND", "合同文件不存在。")
    contract_file, file_object = source
    if file_object.scan_status != "clean" or file_object.storage_status != "stored":
        raise DocumentParseError("DOCUMENT_FILE_NOT_READY", "合同文件尚未准备好解析。")
    organization = session.get(Organization, organization_id)
    if organization is None:
        raise DocumentParseError("DOCUMENT_NOT_FOUND", "文档不存在。")
    settings = organization_settings(organization)
    threshold = (
        ocr_low_confidence_threshold
        if ocr_low_confidence_threshold is not None
        else float(settings["ocr_low_confidence_threshold"])
    )
    page_limit = int(settings["page_limit"])
    parser_name = _parser_name(file_object.media_type, file_object.original_name)
    fingerprint = _fingerprint(
        file_object.sha256,
        contract_file.id,
        parser_name,
        threshold,
        page_limit,
    )
    session.commit()

    parse_failed = False
    with UnitOfWork(session) as unit_of_work:
        document = session.scalar(
            select(DocumentVersion)
            .where(
                DocumentVersion.organization_id == organization_id,
                DocumentVersion.parse_fingerprint == fingerprint,
            )
            .with_for_update()
        )
        if document is None:
            document = DocumentVersion(
                id=uuid4(),
                organization_id=organization_id,
                contract_file_id=contract_file.id,
                parser_name=parser_name,
                parser_version=PARSER_VERSION,
                parse_fingerprint=fingerprint,
                ocr_status="pending",
                status="processing",
            )
            session.add(document)
            session.flush()
        elif document.status == "succeeded":
            unit_of_work.commit()
            return document
        else:
            _clear_document_children(session, document.id, organization_id, file_store)
            document.contract_file_id = contract_file.id
            document.parser_name = parser_name
            document.parser_version = PARSER_VERSION
            document.ocr_status = "pending"
            document.page_count = 0
            document.status = "processing"
            document.text_sha256 = None
            document.error_code = None
            document.error_message = None

        stored_keys: list[str] = []
        try:
            with file_store.open(file_object.storage_key) as source_handle:
                parsed = _parse(
                    source_handle,
                    media_type=file_object.media_type,
                    original_name=file_object.original_name,
                    threshold=threshold,
                    page_limit=page_limit,
                    ocr_engine=ocr_engine,
                )
            _persist_document(
                session,
                document=document,
                parsed=parsed,
                organization_id=organization_id,
                source_file=file_object,
                file_store=file_store,
                stored_keys=stored_keys,
            )
            document.status = "succeeded"
            document.error_code = None
            document.error_message = None
            unit_of_work.commit()
            return document
        except DocumentParseError as exc:
            for storage_key in stored_keys:
                file_store.delete(storage_key)
            document.status = "failed"
            document.ocr_status = "failed"
            document.error_code = exc.code
            document.error_message = exc.message
            unit_of_work.commit()
            return document
        except Exception:
            for storage_key in stored_keys:
                file_store.delete(storage_key)
            unit_of_work.rollback()
            parse_failed = True

    if parse_failed:
        return _record_failed_document(
            session,
            organization_id=organization_id,
            contract_file_id=contract_file.id,
            parser_name=parser_name,
            parse_fingerprint=fingerprint,
            error_code="DOCUMENT_PARSE_FAILED",
            error_message="文档解析失败，请重试或联系管理员。",
        )
    raise RuntimeError("document parsing exited without a result")


def _record_failed_document(
    session: Session,
    *,
    organization_id: UUID,
    contract_file_id: UUID,
    parser_name: str,
    parse_fingerprint: str,
    error_code: str,
    error_message: str,
) -> DocumentVersion:
    session.rollback()
    with UnitOfWork(session) as unit_of_work:
        document = session.scalar(
            select(DocumentVersion)
            .where(
                DocumentVersion.organization_id == organization_id,
                DocumentVersion.parse_fingerprint == parse_fingerprint,
            )
            .with_for_update()
        )
        if document is None:
            document = DocumentVersion(
                id=uuid4(),
                organization_id=organization_id,
                contract_file_id=contract_file_id,
                parser_name=parser_name,
                parser_version=PARSER_VERSION,
                parse_fingerprint=parse_fingerprint,
                ocr_status="failed",
                status="failed",
                error_code=error_code,
                error_message=error_message,
            )
            session.add(document)
        else:
            document.contract_file_id = contract_file_id
            document.parser_name = parser_name
            document.parser_version = PARSER_VERSION
            document.page_count = 0
            document.text_sha256 = None
            document.ocr_status = "failed"
            document.status = "failed"
            document.error_code = error_code
            document.error_message = error_message
        unit_of_work.commit()
        return document


def _parser_name(media_type: str, original_name: str) -> str:
    if media_type == "application/pdf":
        return "pdf"
    if media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return "docx"
    if media_type.startswith("image/"):
        return "image"
    raise DocumentParseError("DOCUMENT_FILE_UNSUPPORTED", "文件类型不支持文档解析。")


def _fingerprint(
    file_sha256: str,
    contract_file_id: UUID,
    parser_name: str,
    threshold: float,
    page_limit: int,
) -> str:
    payload = json.dumps(
        {
            "file_sha256": file_sha256,
            "contract_file_id": str(contract_file_id),
            "parser_name": parser_name,
            "parser_version": PARSER_VERSION,
            "ocr_low_confidence_threshold": threshold,
            "page_limit": page_limit,
        },
        sort_keys=True,
        separators=(",", ":"),
    )
    return hashlib.sha256(payload.encode()).hexdigest()


def _clear_document_children(
    session: Session,
    document_id: UUID,
    organization_id: UUID,
    file_store: LocalFileStore,
) -> None:
    image_files = session.execute(
        select(FileObject.id, FileObject.storage_key)
        .join(
            DocumentPage,
            (DocumentPage.organization_id == FileObject.organization_id)
            & (DocumentPage.image_file_id == FileObject.id),
        )
        .where(
            DocumentPage.organization_id == organization_id,
            DocumentPage.document_version_id == document_id,
        )
    ).all()
    session.execute(
        delete(SourceSpan).where(
            SourceSpan.organization_id == organization_id,
            SourceSpan.document_version_id == document_id,
        )
    )
    session.execute(
        delete(DocumentBlock).where(
            DocumentBlock.organization_id == organization_id,
            DocumentBlock.document_version_id == document_id,
        )
    )
    session.execute(
        delete(DocumentPage).where(
            DocumentPage.organization_id == organization_id,
            DocumentPage.document_version_id == document_id,
        )
    )
    for file_id, storage_key in image_files:
        file_store.delete(storage_key)
        session.execute(
            delete(FileObject).where(
                FileObject.organization_id == organization_id,
                FileObject.id == file_id,
            )
        )


def _parse(
    source: BinaryIO,
    *,
    media_type: str,
    original_name: str,
    threshold: float,
    ocr_engine: OcrEngine | None,
    page_limit: int = DEFAULT_PAGE_LIMIT,
) -> ParsedDocument:
    data = source.read()
    if not data:
        raise DocumentParseError("DOCUMENT_EMPTY", "文档没有可解析内容。")
    if media_type == "application/pdf":
        return _parse_pdf(
            data,
            threshold=threshold,
            page_limit=page_limit,
            ocr_engine=ocr_engine,
        )
    if media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _parse_docx(data)
    if media_type.startswith("image/"):
        return _parse_image(data, threshold=threshold, ocr_engine=ocr_engine)
    raise DocumentParseError(
        "DOCUMENT_FILE_UNSUPPORTED",
        f"不支持解析文件：{Path(original_name).suffix}。",
    )


def _parse_docx(data: bytes) -> ParsedDocument:
    try:
        document = open_docx(io.BytesIO(data))
    except Exception as exc:
        raise DocumentParseError("DOCUMENT_CORRUPTED", "DOCX 文件损坏或无法读取。") from exc
    blocks: list[ParsedBlock] = []
    paragraph_no = 0
    table_no = 0
    for item in _iter_block_items(document):
        if isinstance(item, Paragraph):
            paragraph_no += 1
            text = item.text.strip()
            if text:
                style_name = item.style.name if item.style is not None else ""
                block_type = "heading" if style_name.lower().startswith("heading") else "paragraph"
                blocks.append(
                    ParsedBlock(text=text, block_type=block_type, paragraph_no=paragraph_no)
                )
        else:
            table_no += 1
            table = item
            for row_no, row in enumerate(table.rows, start=1):
                for cell_no, cell in enumerate(row.cells, start=1):
                    text = "\n".join(
                        paragraph.text.strip()
                        for paragraph in cell.paragraphs
                        if paragraph.text.strip()
                    )
                    if text:
                        blocks.append(
                            ParsedBlock(
                                text=text,
                                block_type="table_cell",
                                table_path=f"table[{table_no}]/row[{row_no}]/cell[{cell_no}]",
                            )
                        )
    if not blocks:
        raise DocumentParseError("DOCUMENT_NO_TEXT", "DOCX 中没有可识别的文本或表格。")
    text = "\n".join(block.text for block in blocks)
    return ParsedDocument(
        parser_name="docx",
        page_count=0,
        ocr_status="not_required",
        pages=[],
        blocks=blocks,
        text=text,
    )


def _iter_block_items(parent: DocumentObject | _Cell) -> list[Paragraph | Table]:
    parent_element = parent.element.body if isinstance(parent, DocumentObject) else parent._tc
    items: list[Paragraph | Table] = []
    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            items.append(Paragraph(child, parent))
        elif isinstance(child, CT_Tbl):
            items.append(Table(child, parent))
    return items


def _parse_pdf(
    data: bytes,
    *,
    threshold: float,
    page_limit: int,
    ocr_engine: OcrEngine | None,
) -> ParsedDocument:
    try:
        reader = PdfReader(io.BytesIO(data))
    except (PdfReadError, ValueError, OSError) as exc:
        raise DocumentParseError("DOCUMENT_CORRUPTED", "PDF 文件损坏或无法读取。") from exc
    if reader.is_encrypted:
        raise DocumentParseError("DOCUMENT_ENCRYPTED", "加密 PDF 不能在当前环境中解析。")
    if len(reader.pages) > page_limit:
        raise DocumentParseError(
            "DOCUMENT_PAGE_LIMIT_EXCEEDED",
            "文档页数超过组织配置上限。",
        )
    try:
        pdf = pdfium.PdfDocument(data)
    except Exception as exc:
        raise DocumentParseError("DOCUMENT_RENDER_FAILED", "PDF 页面无法渲染预览。") from exc
    pages: list[ParsedPage] = []
    all_blocks: list[ParsedBlock] = []
    ocr_statuses: list[str] = []
    try:
        for index, pdf_page in enumerate(reader.pages):
            page_no = index + 1
            try:
                extracted = (pdf_page.extract_text() or "").strip()
            except Exception as exc:
                raise DocumentParseError(
                    "DOCUMENT_TEXT_EXTRACTION_FAILED",
                    "PDF 文本提取失败。",
                ) from exc
            try:
                page_size = pdf_page.mediabox
                width: float | None = float(page_size.width)
                height: float | None = float(page_size.height)
            except Exception:
                width = height = None
            image_bytes = _render_pdf_page(pdf, index)
            if extracted:
                blocks = _text_blocks(extracted)
                page = ParsedPage(
                    page_no=page_no,
                    width=width,
                    height=height,
                    text="\n".join(block.text for block in blocks),
                    blocks=blocks,
                    image_bytes=image_bytes,
                    ocr_status="not_required",
                    ocr_confidence=None,
                )
            else:
                page = _ocr_page(
                    page_no=page_no,
                    image_bytes=image_bytes,
                    threshold=threshold,
                    ocr_engine=ocr_engine,
                    width=width,
                    height=height,
                )
            pages.append(page)
            all_blocks.extend(page.blocks)
            ocr_statuses.append(page.ocr_status)
    finally:
        pdf.close()
    return ParsedDocument(
        parser_name="pdf",
        page_count=len(pages),
        ocr_status=_document_ocr_status(ocr_statuses),
        pages=pages,
        blocks=all_blocks,
        text="\n".join(page.text for page in pages if page.text),
    )


def _render_pdf_page(pdf: Any, index: int) -> bytes:
    try:
        rendered = pdf[index].render(scale=1.5)
        with rendered.to_pil() as image:
            output = io.BytesIO()
            image.save(output, format="PNG")
            return output.getvalue()
    except Exception as exc:
        raise DocumentParseError("DOCUMENT_RENDER_FAILED", "PDF 页面无法渲染预览。") from exc


def _parse_image(data: bytes, *, threshold: float, ocr_engine: OcrEngine | None) -> ParsedDocument:
    try:
        with Image.open(io.BytesIO(data)) as image:
            image.load()
            width, height = image.size
    except Exception as exc:
        raise DocumentParseError("DOCUMENT_CORRUPTED", "图片文件损坏或无法读取。") from exc
    page = _ocr_page(
        page_no=1,
        image_bytes=data,
        threshold=threshold,
        ocr_engine=ocr_engine,
        width=float(width),
        height=float(height),
    )
    return ParsedDocument(
        parser_name="image",
        page_count=1,
        ocr_status=_document_ocr_status([page.ocr_status]),
        pages=[page],
        blocks=page.blocks,
        text=page.text,
    )


def _text_blocks(text: str) -> list[ParsedBlock]:
    return [
        ParsedBlock(text=line.strip(), block_type="paragraph")
        for line in text.splitlines()
        if line.strip()
    ]


def _ocr_page(
    *,
    page_no: int,
    image_bytes: bytes,
    threshold: float,
    ocr_engine: OcrEngine | None,
    width: float | None,
    height: float | None,
) -> ParsedPage:
    engine = ocr_engine or PaddleOcrEngine()
    try:
        lines = engine.recognize(image_bytes)
    except Exception:
        return ParsedPage(
            page_no=page_no,
            width=width,
            height=height,
            text="",
            blocks=[],
            image_bytes=image_bytes,
            ocr_status="failed",
            ocr_confidence=None,
            error_code="OCR_FAILED",
            error_message="页面 OCR 失败，请人工核对原图。",
        )
    if not lines:
        return ParsedPage(
            page_no=page_no,
            width=width,
            height=height,
            text="",
            blocks=[],
            image_bytes=image_bytes,
            ocr_status="blank",
            ocr_confidence=0.0,
            error_code="OCR_BLANK_PAGE",
            error_message="页面未识别到可用文字，请人工核对原图。",
        )
    blocks = [ParsedBlock(text=line.text, block_type="ocr_line", bbox=line.bbox) for line in lines]
    confidence = sum(line.confidence for line in lines) / len(lines)
    status = "completed" if confidence >= threshold else "low_confidence"
    return ParsedPage(
        page_no=page_no,
        width=width,
        height=height,
        text="\n".join(block.text for block in blocks),
        blocks=blocks,
        image_bytes=image_bytes,
        ocr_status=status,
        ocr_confidence=confidence,
        error_code="OCR_LOW_CONFIDENCE" if status == "low_confidence" else None,
        error_message="页面 OCR 置信度较低，请人工复核。" if status == "low_confidence" else None,
    )


def _document_ocr_status(statuses: list[str]) -> str:
    if not statuses or all(status == "not_required" for status in statuses):
        return "not_required"
    if any(status == "failed" for status in statuses):
        return (
            "partial"
            if any(status in {"completed", "low_confidence", "not_required"} for status in statuses)
            else "failed"
        )
    if any(status in {"low_confidence", "blank"} for status in statuses):
        return "low_confidence"
    return "completed"


def _persist_document(
    session: Session,
    *,
    document: DocumentVersion,
    parsed: ParsedDocument,
    organization_id: UUID,
    source_file: FileObject,
    file_store: LocalFileStore,
    stored_keys: list[str],
) -> None:
    block_rows: list[tuple[DocumentBlock, ParsedBlock, UUID | None]] = []
    for parsed_page in parsed.pages:
        observe_ocr_page(parsed_page.ocr_status)
        image_file_id = source_file.id if parsed.parser_name == "image" else _store_page_image(
            session,
            organization_id=organization_id,
            document_id=document.id,
            page_no=parsed_page.page_no,
            image_bytes=parsed_page.image_bytes,
            file_store=file_store,
            stored_keys=stored_keys,
        )
        page = DocumentPage(
            id=uuid4(),
            organization_id=organization_id,
            document_version_id=document.id,
            page_no=parsed_page.page_no,
            width=parsed_page.width,
            height=parsed_page.height,
            text=parsed_page.text,
            image_file_id=image_file_id,
            ocr_status=parsed_page.ocr_status,
            ocr_confidence=parsed_page.ocr_confidence,
            error_code=parsed_page.error_code,
            error_message=parsed_page.error_message,
        )
        session.add(page)
        for parsed_block in parsed_page.blocks:
            block_rows.append((
                DocumentBlock(
                    id=uuid4(),
                    organization_id=organization_id,
                    document_version_id=document.id,
                    page_id=page.id,
                    order_no=len(block_rows) + 1,
                    block_type=parsed_block.block_type,
                    paragraph_no=parsed_block.paragraph_no,
                    table_path=parsed_block.table_path,
                    text=parsed_block.text,
                    bbox_json=parsed_block.bbox,
                ),
                parsed_block,
                page.id,
            ))
    session.flush()
    for parsed_block in parsed.blocks:
        if parsed.parser_name == "docx":
            block_rows.append((
                DocumentBlock(
                    id=uuid4(),
                    organization_id=organization_id,
                    document_version_id=document.id,
                    page_id=None,
                    order_no=len(block_rows) + 1,
                    block_type=parsed_block.block_type,
                    paragraph_no=parsed_block.paragraph_no,
                    table_path=parsed_block.table_path,
                    text=parsed_block.text,
                    bbox_json=parsed_block.bbox,
                ),
                parsed_block,
                None,
            ))
    for block, _, _ in block_rows:
        session.add(block)
    session.flush()
    for block, parsed_block, page_id in block_rows:
        session.add(
            SourceSpan(
                id=uuid4(),
                organization_id=organization_id,
                document_version_id=document.id,
                page_id=page_id,
                block_id=block.id,
                start_offset=0,
                end_offset=len(parsed_block.text),
                bbox_json=parsed_block.bbox,
                quote=parsed_block.text,
                quote_sha256=hashlib.sha256(parsed_block.text.encode()).hexdigest(),
            )
        )
    document.page_count = parsed.page_count
    document.ocr_status = parsed.ocr_status
    document.text_sha256 = hashlib.sha256(parsed.text.encode()).hexdigest()
    document.status = "succeeded"
    session.flush()


def _store_page_image(
    session: Session,
    *,
    organization_id: UUID,
    document_id: UUID,
    page_no: int,
    image_bytes: bytes,
    file_store: LocalFileStore,
    stored_keys: list[str],
) -> UUID:
    file_id = uuid4()
    storage_key = f"org/{organization_id}/documents/{document_id}/pages/{file_id}"
    write_operation_id = create_file_write_journal(
        session,
        organization_id=organization_id,
        storage_key=storage_key,
    )
    size_bytes, sha256 = file_store.put(io.BytesIO(image_bytes), storage_key)
    stored_keys.append(storage_key)
    file_object = FileObject(
        id=file_id,
        organization_id=organization_id,
        storage_key=storage_key,
        original_name=f"document-page-{page_no}.png",
        media_type="image/png",
        size_bytes=size_bytes,
        sha256=sha256,
        scan_status="clean",
        storage_status="stored",
    )
    session.add(file_object)
    session.flush()
    finalize_file_write_journal(
        session,
        operation_id=write_operation_id,
        file_object_id=file_object.id,
    )
    return file_id


def document_kind(document: DocumentVersion) -> str:
    return document.parser_name


def source_kind(document: DocumentVersion, block: DocumentBlock) -> str:
    if block.table_path:
        return "docx_table_cell"
    if block.page_id is None:
        return "docx_paragraph"
    return "image_page" if document.parser_name == "image" else "pdf_page"
