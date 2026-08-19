from io import BytesIO
from types import SimpleNamespace

from docx import Document
from PIL import Image
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from backend.app.integrations.ocr import OcrLine
from backend.app.integrations.ocr.paddleocr import _parse_predict_results
from backend.app.modules.documents.service import DocumentParseError, _parse


class FakeOcr:
    def __init__(self, lines: list[OcrLine] | None = None, error: bool = False) -> None:
        self.lines = lines or []
        self.error = error

    def recognize(self, image: bytes) -> list[OcrLine]:
        if self.error:
            raise RuntimeError("fake OCR failure")
        return self.lines


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("采购协议", level=1)
    document.add_paragraph("第一段保留原始顺序。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "甲方"
    table.cell(0, 1).text = "乙方"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _text_pdf_bytes() -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=300, height=400)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font})}
    )
    content = DecodedStreamObject()
    content.set_data(b"BT /F1 12 Tf 20 350 Td (Text PDF evidence) Tj ET")
    page[NameObject("/Contents")] = content
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def _image_bytes() -> bytes:
    output = BytesIO()
    Image.new("RGB", (120, 80), "white").save(output, format="PNG")
    return output.getvalue()


def test_docx_preserves_paragraph_table_order_without_page_numbers() -> None:
    parsed = _parse(
        BytesIO(_docx_bytes()),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        original_name="sample.docx",
        threshold=0.8,
        ocr_engine=FakeOcr(),
    )

    assert parsed.page_count == 0
    assert [block.block_type for block in parsed.blocks] == [
        "heading",
        "paragraph",
        "table_cell",
        "table_cell",
    ]
    assert parsed.blocks[2].table_path == "table[1]/row[1]/cell[1]"
    assert all(block.paragraph_no is None or block.paragraph_no > 0 for block in parsed.blocks)


def test_text_pdf_preserves_page_and_text_blocks() -> None:
    parsed = _parse(
        BytesIO(_text_pdf_bytes()),
        media_type="application/pdf",
        original_name="sample.pdf",
        threshold=0.8,
        ocr_engine=FakeOcr(error=True),
    )

    assert parsed.page_count == 1
    assert parsed.pages[0].ocr_status == "not_required"
    assert parsed.pages[0].text == "Text PDF evidence"
    assert parsed.pages[0].blocks[0].text == "Text PDF evidence"
    assert parsed.pages[0].image_bytes.startswith(b"\x89PNG")


def test_pdf_page_limit_is_rejected_before_rendering() -> None:
    try:
        _parse(
            BytesIO(_text_pdf_bytes()),
            media_type="application/pdf",
            original_name="sample.pdf",
            threshold=0.8,
            page_limit=0,
            ocr_engine=FakeOcr(error=True),
        )
    except DocumentParseError as exc:
        assert exc.code == "DOCUMENT_PAGE_LIMIT_EXCEEDED"
    else:
        raise AssertionError("expected page limit error")


def test_scanned_pdf_records_ocr_bbox_and_low_confidence() -> None:
    parsed = _parse(
        BytesIO(_blank_pdf_bytes()),
        media_type="application/pdf",
        original_name="scan.pdf",
        threshold=0.8,
        ocr_engine=FakeOcr(
            [OcrLine("扫描文本", 0.65, {"x": 10.0, "y": 20.0, "width": 80.0, "height": 18.0})]
        ),
    )

    assert parsed.pages[0].ocr_status == "low_confidence"
    assert parsed.pages[0].blocks[0].bbox == {
        "x": 10.0,
        "y": 20.0,
        "width": 80.0,
        "height": 18.0,
    }


def _blank_pdf_bytes() -> bytes:
    writer = PdfWriter()
    writer.add_blank_page(width=300, height=400)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def test_image_ocr_failure_is_explicit() -> None:
    parsed = _parse(
        BytesIO(_image_bytes()),
        media_type="image/png",
        original_name="sample.png",
        threshold=0.8,
        ocr_engine=FakeOcr(error=True),
    )

    assert parsed.pages[0].ocr_status == "failed"
    assert parsed.pages[0].error_code == "OCR_FAILED"
    assert parsed.pages[0].text == ""


def test_paddle_predict_result_preserves_xyxy_bbox() -> None:
    result = SimpleNamespace(
        json={
            "res": {
                "rec_texts": ["证据"],
                "rec_scores": [0.91],
                "rec_boxes": [[10, 20, 80, 48]],
            }
        }
    )

    lines = _parse_predict_results([result])

    assert lines == [OcrLine("证据", 0.91, {"x": 10.0, "y": 20.0, "width": 70.0, "height": 28.0})]
